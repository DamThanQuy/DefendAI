"""
Unified DOCX Extractor — Trích xuất toàn diện từ file DOCX.

Kết hợp:
- Native text extraction (python-docx): paragraphs, tables, headers, footers, TOC structure
- Image extraction (zipfile): embedded images từ word/media/
- Vision-ready output: base64 data URIs cho vision reader

Dùng cho cả document_parser.py (RAG) và deliverable_classify.py (classification).
"""

from __future__ import annotations

import base64
import logging
import zipfile
from dataclasses import dataclass, field
from io import BytesIO
from typing import List, Optional

from docx import Document as DocxDocument
from PIL import Image

logger = logging.getLogger(__name__)


@dataclass
class DocxExtractResult:
    """Kết quả extract đầy đủ từ DOCX."""
    # Native text content
    text: str = ""
    # Structured headings for RAG chunking
    headings: List[dict] = field(default_factory=list)
    # Header/footer text (metadata: tên SV, mã đồ án, ngày...)
    headers: List[str] = field(default_factory=list)
    footers: List[str] = field(default_factory=list)
    # Tables as structured data
    tables: List[List[List[str]]] = field(default_factory=list)
    # Embedded images as base64 data URIs (cho vision reader)
    images: List[str] = field(default_factory=list)
    # Image metadata
    image_metadata: List[dict] = field(default_factory=list)
    # Raw paragraphs for fallback
    paragraphs: List[str] = field(default_factory=list)


# ---------------------------------------------------------------------------
# Constants
# ---------------------------------------------------------------------------
# Skip tiny images (icons, bullets, decorative)
MIN_IMAGE_WIDTH = 200
MIN_IMAGE_HEIGHT = 200
# Max images to send to vision reader (cost/latency control)
MAX_VISION_IMAGES = 4
# Supported image MIME types
IMAGE_MIME_TYPES = {
    "png": "image/png",
    "jpg": "image/jpeg",
    "jpeg": "image/jpeg",
    "gif": "image/gif",
    "webp": "image/webp",
    "bmp": "image/bmp",
}


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------
def extract_docx_full(file_bytes: bytes) -> DocxExtractResult:
    """
    Extract everything from DOCX: text, structure, headers/footers, images.
    
    Args:
        file_bytes: Raw DOCX file bytes
        
    Returns:
        DocxExtractResult with all extracted content
    """
    result = DocxExtractResult()
    
    try:
        # 1. Native text extraction with python-docx
        doc = DocxDocument(BytesIO(file_bytes))
        _extract_native_content(doc, result)
        
        # 2. Image extraction from word/media/
        _extract_images(file_bytes, result)
        
        # 3. Build combined text (for backward compatibility)
        result.text = _build_combined_text(result)
        
        logger.info(
            "DOCX extracted: %d chars, %d headings, %d tables, %d images",
            len(result.text), len(result.headings), len(result.tables), len(result.images)
        )
        
    except Exception as exc:
        logger.exception("DOCX extraction failed: %s", exc)
        result.text = f"[DOCX extraction failed: {exc}]"
    
    return result


def extract_docx_for_vision(file_bytes: bytes, max_images: int = MAX_VISION_IMAGES) -> tuple[str, List[str]]:
    """
    Extract text + images optimized for vision reader (Option B pattern).
    
    Returns:
        (native_text, list_of_base64_data_uris)
    """
    result = extract_docx_full(file_bytes)
    # Cap images for vision reader
    vision_images = result.images[:max_images]
    return result.text, vision_images


# ---------------------------------------------------------------------------
# Internal: Native content extraction
# ---------------------------------------------------------------------------
def _extract_native_content(doc: DocxDocument, result: DocxExtractResult) -> None:
    """Extract text, headings, headers, footers, tables from python-docx."""
    
    # Headers
    for section in doc.sections:
        for para in section.header.paragraphs:
            if para.text.strip():
                result.headers.append(f"[HEADER] {para.text.strip()}")
        for para in section.footer.paragraphs:
            if para.text.strip():
                result.footers.append(f"[FOOTER] {para.text.strip()}")
    
    # Paragraphs with heading detection
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        
        result.paragraphs.append(text)
        
        # Detect heading style
        style_name = para.style.name.lower() if para.style else ""
        if style_name.startswith("heading"):
            try:
                level = int(style_name.replace("heading", "").strip() or "1")
            except ValueError:
                level = 1
            result.headings.append({
                "level": level,
                "text": text,
                "style": style_name
            })
    
    # Tables
    for table in doc.tables:
        table_data = []
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells]
            if any(row_data):
                table_data.append(row_data)
        if table_data:
            result.tables.append(table_data)


def _extract_images(file_bytes: bytes, result: DocxExtractResult) -> None:
    """Extract embedded images from word/media/ as base64 data URIs."""
    
    try:
        with zipfile.ZipFile(BytesIO(file_bytes)) as archive:
            for info in archive.infolist():
                if info.is_dir() or not info.filename.startswith("word/media/"):
                    continue
                
                ext = info.filename.rsplit(".", 1)[-1].lower()
                mime = IMAGE_MIME_TYPES.get(ext)
                if not mime:
                    continue
                
                raw_img = archive.read(info)
                
                # Filter tiny images (icons, bullets)
                try:
                    with Image.open(BytesIO(raw_img)) as img:
                        w, h = img.size
                        if w < MIN_IMAGE_WIDTH and h < MIN_IMAGE_HEIGHT:
                            logger.debug("Skipping small image %s: %dx%d", info.filename, w, h)
                            continue
                        # Store metadata
                        result.image_metadata.append({
                            "filename": info.filename,
                            "width": w,
                            "height": h,
                            "mime": mime,
                            "size_bytes": len(raw_img)
                        })
                except Exception:
                    # If can't open, still include (might be valid)
                    pass
                
                # Convert to base64 data URI
                b64 = base64.b64encode(raw_img).decode()
                data_uri = f"data:{mime};base64,{b64}"
                result.images.append(data_uri)
                
    except Exception as exc:
        logger.debug("DOCX image extraction failed: %s", exc)


def _build_combined_text(result: DocxExtractResult) -> str:
    """Build combined text from all extracted parts (for backward compat)."""
    parts = []
    
    # Headers first (metadata)
    if result.headers:
        parts.extend(result.headers)
    
    # Headings structure
    if result.headings:
        parts.append("=== DOCUMENT STRUCTURE ===")
        for h in result.headings:
            indent = "  " * (h["level"] - 1)
            parts.append(f"{indent}{h['text']}")
    
    # Paragraphs
    if result.paragraphs:
        parts.append("=== CONTENT ===")
        parts.extend(result.paragraphs)
    
    # Tables
    if result.tables:
        parts.append("=== TABLES ===")
        for i, table in enumerate(result.tables):
            parts.append(f"--- Table {i+1} ---")
            for row in table:
                parts.append("\t".join(row))
    
    # Footers
    if result.footers:
        parts.extend(result.footers)
    
    return "\n".join(parts)


# ---------------------------------------------------------------------------
# Backward compatibility helpers
# ---------------------------------------------------------------------------
def extract_docx_text_only(file_bytes: bytes) -> str:
    """Backward compatible: extract text only (like old _extract_docx)."""
    result = extract_docx_full(file_bytes)
    return result.text


def extract_docx_text_and_images(file_bytes: bytes, max_chars: int = 100000) -> tuple[str, List[str]]:
    """Backward compatible: extract text preview + images (like _extract_text_preview)."""
    result = extract_docx_full(file_bytes)
    text = result.text[:max_chars]
    if len(result.text) > max_chars:
        text += f"\n\n[...{len(result.text) - max_chars} chars more truncated...]"
    return text, result.images